from groq import Groq
from langchain_groq import ChatGroq
import streamlit as st
import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from io import BytesIO
from langchain.vectorstores import Chroma
from langchain.embeddings import HuggingFaceEmbeddings

# Verificar que la clave de API se ha cargado
api_key= st.secrets.API_KEY
# Inicializar el cliente de Groq
client = Groq(api_key=api_key)

# Inicializar el modelo de chat
chat_groq = ChatGroq(model="llama-3.1-70b-versatile",api_key=api_key)
# Configurar la página
st.set_page_config(
    page_title="Microcurriculum UdeA con IA",
    page_icon="🩺",
    layout="centered",
    initial_sidebar_state="auto",
)

st.sidebar.image("UdeA+simplificado-01.png", width=200)
st.sidebar.title("Microcurriculum con IA")
st.sidebar.write("Entrega Diplomado Pedagogía UdeA 2024-2")
st.sidebar.write("Alejandro Hernández-Arango, MD, MSc")

st.title("PROGRAMA OFICIAL DE CURSO")
st.write("Aplica para Pregrado y Posgrado")

def mejorar_texto_con_IA(var_name, var_value, vectordb_list, nombre_curso):
    context_docs = []
    search=var_name+"\n"+var_value
    for vectordb in vectordb_list:
        docs = vectordb.similarity_search(search, k=3)
        context_docs.extend([doc.page_content for doc in docs])
    context = "\n".join(context_docs)
    prompt = f"""Tu eres una inteligencia artificial experta en pedagogia. Tu tarea es mejorar el siguiente microcurriculum del curso {nombre_curso} texto en la sección {var_name} teniendo en cuenta el contexto proporcionado. Solo dame el texto mejorado

                Contexto:         
                {context}

                Texto:
                {var_value}

                Texto mejorado:"""
    improved_text = chat_groq.invoke(prompt)
    improved_text = improved_text.content
    return improved_text, context


# Función para generar el documento
# Modifica la función 'crear_documento' para aceptar 'variables' como parámetro:
def crear_documento(nombre_curso, variables):
    doc = Document()
    
    # Insertar la imagen
    doc.add_picture("UdeA+simplificado-01.png", width=Inches(2))

    # Formato de Títulos
    title_format = doc.add_paragraph().add_run("PROGRAMA OFICIAL DE CURSO\n(Pregrado y Posgrado)\nUNIVERSIDAD DE ANTIOQUIA")
    title_format.bold = True
    title_format.font.size = Pt(14)

    # Información General
    doc.add_paragraph("\nINFORMACIÓN GENERAL")
    doc.add_paragraph(f"Nombre del curso: {nombre_curso}")
    doc.add_paragraph(f"Programa académico al que pertenece: {variables.get('programa_academico', '')}")
    doc.add_paragraph(f"Unidad académica: {variables.get('unidad_academica', '')}")
    doc.add_paragraph(f"Programa(s) en los cuales se ofrece el curso: {variables.get('programas_ofrece', '')}")
    doc.add_paragraph(f"Vigencia: {variables.get('vigencia', '')}")
    doc.add_paragraph(f"Código curso: {variables.get('codigo_curso', '')}")
    tipo_curso = variables.get('tipo_curso', '')
    tipo_curso_otro = variables.get('tipo_curso_otro', '')
    doc.add_paragraph(f"Tipo de curso: {tipo_curso if tipo_curso != 'Otro' else tipo_curso_otro}")
    caracteristicas_curso = variables.get('caracteristicas_curso', [])
    doc.add_paragraph(f"Características del curso: {', '.join(caracteristicas_curso)}")
    modalidad = variables.get('modalidad', '')
    modalidad_otra = variables.get('modalidad_otra', '')
    doc.add_paragraph(f"Modalidad educativa del curso: {modalidad if modalidad != 'Otra' else modalidad_otra}")
    doc.add_paragraph(f"Nombre del área, núcleo o componente: {variables.get('nombre_area', '')}")
    doc.add_paragraph(f"Prerrequisitos: {variables.get('prerrequisitos', '')}")
    doc.add_paragraph(f"Correquisitos: {variables.get('correquisitos', '')}")
    doc.add_paragraph(f"Número de créditos académicos: {variables.get('num_creditos', '')}")
    doc.add_paragraph(f"Horas totales de interacción estudiante-profesor: {variables.get('horas_interaccion', '')}")
    doc.add_paragraph(f"Horas totales de trabajo independiente: {variables.get('horas_independiente', '')}")
    doc.add_paragraph(f"Horas totales del curso: {variables.get('horas_totales', '')}")
    doc.add_paragraph(f"Horas totales de actividades académicas teóricas: {variables.get('horas_teoricas', '')}")
    doc.add_paragraph(f"Horas totales de actividades académicas prácticas: {variables.get('horas_practicas', '')}")
    doc.add_paragraph(f"Horas totales de actividades académicas teórico-prácticas: {variables.get('horas_teorico_practicas', '')}")

    # Relaciones con el perfil
    doc.add_paragraph("\nRELACIONES CON EL PERFIL")
    doc.add_paragraph(variables.get('relaciones_perfil', ''))

    # Intencionalidades formativas
    doc.add_paragraph("\nINTENCIONALIDADES FORMATIVAS")
    doc.add_paragraph(variables.get('intencionalidades', ''))

    # Aportes del curso
    doc.add_paragraph("\nAPORTES DEL CURSO A LA FORMACIÓN INTEGRAL Y A LA FORMACIÓN EN INVESTIGACIÓN")
    doc.add_paragraph(variables.get('aportes_formacion', ''))

    # Líneas de sentido y derrotero de preguntas
    doc.add_paragraph("\nLÍNEAS DE SENTIDO Y DERROTERO DE PREGUNTAS")
    doc.add_paragraph(f"Líneas de sentido: {variables.get('lineas_sentido', '')}")
    doc.add_paragraph(f"Derrotero de preguntas: {variables.get('derrotero_preguntas', '')}")

    # Descripción de conocimientos
    doc.add_paragraph("\nDESCRIPCIÓN DE LOS CONOCIMIENTOS Y/O SABERES")
    doc.add_paragraph(variables.get('descripcion_conocimientos', ''))

    # Metodología
    doc.add_paragraph("\nMETODOLOGÍA")
    doc.add_paragraph(f"Estrategias didácticas: {variables.get('estrategias_didacticas', '')}")
    doc.add_paragraph(f"Metodología: {variables.get('metodologia', '')}")
    doc.add_paragraph(f"Medios y recursos didácticos: {variables.get('medios_recursos', '')}")
    doc.add_paragraph(f"Formas de interacción: {variables.get('formas_interaccion', '')}")
    doc.add_paragraph(f"Estrategias de internacionalización: {variables.get('estrategias_internacionalizacion', '')}")
    doc.add_paragraph(f"Estrategias de diversidad: {variables.get('estrategias_diversidad', '')}")

    # Evaluación
    doc.add_paragraph("\nEVALUACIÓN")
    doc.add_paragraph(f"Concepción de evaluación y modalidades: {variables.get('concepcion_evaluacion', '')}")
    doc.add_paragraph(f"Procesos de aprendizaje abordados: {variables.get('procesos_aprendizaje', '')}")
    doc.add_paragraph("Momentos de evaluación y porcentajes:")
    doc.add_paragraph(variables.get('momentos_evaluacion', ''))

    # Bibliografía
    doc.add_paragraph("\nBIBLIOGRAFÍA Y OTRAS FUENTES")
    doc.add_paragraph(variables.get('bibliografia', ''))

    # Comunidad académica
    doc.add_paragraph("\nCOMUNIDAD ACADÉMICA QUE PARTICIPÓ EN LA ELABORACIÓN DEL MICROCURRÍCULO")
    doc.add_paragraph(f"Nombres y apellidos: {variables.get('nombres_apellidos', '')}")
    doc.add_paragraph(f"Unidad académica: {variables.get('unidad_academica_participantes', '')}")
    doc.add_paragraph(f"Formación académica: {variables.get('formacion_academica', '')}")

    # Sección de aprobación del Consejo de Unidad Académica
    doc.add_paragraph("\nAPROBACIÓN DEL CONSEJO DE UNIDAD ACADÉMICA").bold = True

    aprobacion_parrafo = doc.add_paragraph()
    aprobacion_parrafo.add_run("Aprobado en Acta número ").bold = False
    aprobacion_parrafo.add_run("_______________").underline = True
    aprobacion_parrafo.add_run(" del ").bold = False
    aprobacion_parrafo.add_run("_______________").underline = True

    # Información del Secretario
    doc.add_paragraph("\nNombre completo del Secretario del Consejo de la Unidad Académica:")
    doc.add_paragraph("_________________________").underline = True  # Línea para nombre del Secretario
    
    doc.add_paragraph("Firma:")
    doc.add_paragraph("_________________________").underline = True  # Línea para la firma
    
    doc.add_paragraph("Cargo:")
    doc.add_paragraph("_________________________").underline = True  # Línea para el cargo

    # Guardar el documento en un objeto de bytes
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)

    return doc_buffer




with st.expander("INFORMACIÓN GENERAL"):
        nombre_curso = st.text_input("Nombre del curso:", value=st.session_state.get('nombre_curso', ''), key='nombre_curso')
        programa_academico = st.text_input("Programa académico al que pertenece:", value=st.session_state.get('programa_academico', ''), key='programa_academico')
        unidad_academica = st.text_input("Unidad académica:", value=st.session_state.get('unidad_academica', ''), key='unidad_academica')
        programas_ofrece = st.text_input("Programa(s) académico(s) en los cuales se ofrece el curso:", value=st.session_state.get('programas_ofrece', ''), key='programas_ofrece')
        vigencia = st.text_input("Vigencia (Periodo académico o Cohorte):", value=st.session_state.get('vigencia', ''), key='vigencia')
        codigo_curso = st.text_input("Código curso (en MARES):", value=st.session_state.get('codigo_curso', ''), key='codigo_curso')


        tipo_curso_options = ["Obligatorio", "Electivo", "Otro"]
        tipo_curso_index = tipo_curso_options.index(st.session_state.get('tipo_curso', 'Obligatorio'))
        tipo_curso = st.selectbox("Tipo de curso:", tipo_curso_options, index=tipo_curso_index, key='tipo_curso')
        if tipo_curso == "Otro":
            tipo_curso_otro = st.text_input("En caso de elegir 'Otro', indique cuál:", value=st.session_state.get('tipo_curso_otro', ''), key='tipo_curso_otro')
        else:
            tipo_curso_otro = ''

        caracteristicas_curso = st.multiselect(
            "Características del curso:",
            ["Validable", "Habilitable", "Clasificable", "Evaluación de suficiencia (posgrado)"],
            default=st.session_state.get('caracteristicas_curso', []),
            key='caracteristicas_curso'
        )

        modalidad_options = ["Presencial", "Virtual", "Mixta", "Otra"]
        modalidad_index = modalidad_options.index(st.session_state.get('modalidad', 'Presencial'))
        modalidad = st.selectbox("Modalidad educativa del curso:", modalidad_options, index=modalidad_index, key='modalidad')
        if modalidad == "Otra":
            modalidad_otra = st.text_input("En caso de elegir 'Otra', indique cuál:", value=st.session_state.get('modalidad_otra', ''), key='modalidad_otra')
        else:
            modalidad_otra = ''

        nombre_area = st.text_input("Nombre del área, núcleo o componente de la organización curricular a la que pertenece el curso:", value=st.session_state.get('nombre_area', ''), key='nombre_area')

        prerrequisitos = st.text_area("Prerrequisitos (con nombre y código en MARES):", value=st.session_state.get('prerrequisitos', ''), key='prerrequisitos')
        correquisitos = st.text_area("Correquisitos (con nombre y código en MARES):", value=st.session_state.get('correquisitos', ''), key='correquisitos')

        num_creditos = st.number_input("Número de créditos académicos (Acuerdo Académico 576 de marzo de 2021):", min_value=0, step=1, value=st.session_state.get('num_creditos', 0), key='num_creditos')
        horas_interaccion = st.number_input("Horas totales de interacción estudiante-profesor:", min_value=0, step=1, value=st.session_state.get('horas_interaccion', 0), key='horas_interaccion')
        horas_independiente = st.number_input("Horas totales de trabajo independiente:", min_value=0, step=1, value=st.session_state.get('horas_independiente', 0), key='horas_independiente')
        horas_totales = st.number_input("Horas totales del curso (suma de las horas anteriores):", min_value=0, step=1, value=st.session_state.get('horas_totales', 0), key='horas_totales')
        horas_teoricas = st.number_input("Horas totales de actividades académicas teóricas:", min_value=0, step=1, value=st.session_state.get('horas_teoricas', 0), key='horas_teoricas')
        horas_practicas = st.number_input("Horas totales de actividades académicas prácticas:", min_value=0, step=1, value=st.session_state.get('horas_practicas', 0), key='horas_practicas')
        horas_teorico_practicas = st.number_input("Horas totales de actividades académicas teórico-prácticas:", min_value=0, step=1, value=st.session_state.get('horas_teorico_practicas', 0), key='horas_teorico_practicas')

# For 'relaciones_perfil'
with st.expander("RELACIONES CON EL PERFIL"):
    var_name = 'relaciones_perfil'
    var_value = st.session_state.get('improved_' + var_name, st.session_state.get(var_name, ''))
    relaciones_perfil = st.text_area(
        "Describir el propósito del curso en relación con los perfiles del programa académico.",
        value=var_value,
        key=var_name
    )
    if st.button("Mejorar con IA✨", key='btn_' + var_name):
        with st.spinner(f'Mejorando {var_name} con IA...'):
            try:
                embeddings = HuggingFaceEmbeddings()
                vectordb_PEI = Chroma(persist_directory="./vector_db_PEI", embedding_function=embeddings)
                vectordb_guia = Chroma(persist_directory="./vector_db_guia", embedding_function=embeddings)
                vectordb_list = [vectordb_PEI, vectordb_guia]
                improved_text, context = mejorar_texto_con_IA(var_name, relaciones_perfil, vectordb_list, nombre_curso)
                st.session_state['improved_' + var_name] = improved_text
                st.session_state['context_' + var_name] = context
                st.rerun()
            except Exception as e:
                st.warning(e)
                pass
    if 'context_' + var_name in st.session_state:
        st.warning("Fragmentos del PEI-UdeA usados por la IA")
        st.info(st.session_state['context_' + var_name])
                                 
# For 'intencionalidades'
with st.expander("INTENCIONALIDADES FORMATIVAS"):
    var_name = 'intencionalidades'
    var_value = st.session_state.get('improved_' + var_name, st.session_state.get(var_name, ''))
    intencionalidades = st.text_area(
        "Explicitar los elementos orientadores del curso de acuerdo con el diseño curricular del programa académico.",
        value=var_value,
        key=var_name
    )
    if st.button("Mejorar con IA✨", key='btn_' + var_name):
        with st.spinner(f'Mejorando {var_name} con IA...'):
            try:
                embeddings = HuggingFaceEmbeddings()
                vectordb_PEI = Chroma(persist_directory="./vector_db_PEI", embedding_function=embeddings)
                vectordb_guia = Chroma(persist_directory="./vector_db_guia", embedding_function=embeddings)
                vectordb_list = [vectordb_PEI, vectordb_guia]
                improved_text, context = mejorar_texto_con_IA(var_name, intencionalidades, vectordb_list, nombre_curso)
                st.session_state['improved_' + var_name] = improved_text
                st.session_state['context_' + var_name] = context
                st.rerun()
            except Exception as e:
                st.warning(e)
                pass
    if 'context_' + var_name in st.session_state:
        st.warning("Fragmentos del PEI-UdeA usados por la IA")
        st.info(st.session_state['context_' + var_name])
                                 
# For 'aportes_formacion', 'lineas_sentido', 'derrotero_preguntas'
with st.expander("APORTES DEL CURSO A LA FORMACIÓN INTEGRAL Y A LA FORMACIÓN EN INVESTIGACIÓN"):
    # 'aportes_formacion'
    var_name = 'aportes_formacion'
    var_value = st.session_state.get('improved_' + var_name, st.session_state.get(var_name, ''))
    aportes_formacion = st.text_area(
        "Describir cómo el curso hace aportes a la formación integral y a la formación en investigación.",
        value=var_value,
        key=var_name
    )
    if st.button("Mejorar con IA✨", key='btn_' + var_name):
        with st.spinner(f"Mejorando {var_name} con IA..."):
            try:
                embeddings = HuggingFaceEmbeddings()
                vectordb_PEI = Chroma(persist_directory="./vector_db_PEI", embedding_function=embeddings)
                vectordb_guia = Chroma(persist_directory="./vector_db_guia", embedding_function=embeddings)
                vectordb_list = [vectordb_PEI, vectordb_guia]
                improved_text, context = mejorar_texto_con_IA(var_name, aportes_formacion, vectordb_list, nombre_curso)
                st.session_state['improved_' + var_name] = improved_text
                st.session_state['context_' + var_name] = context
                st.rerun()
            except Exception as e:
                st.warning(e)
                pass
    if 'context_' + var_name in st.session_state:
        st.warning("Fragmentos del PEI-UdeA usados por la IA")
        st.info(st.session_state['context_' + var_name])

    # 'lineas_sentido'
    var_name = 'lineas_sentido'
    var_value = st.session_state.get('improved_' + var_name, st.session_state.get(var_name, ''))
    lineas_sentido = st.text_area(
        "Líneas de sentido",
        value=var_value,
        key=var_name
    )
    if st.button("Mejorar con IA✨", key='btn_' + var_name):
        with st.spinner(f"Mejorando {var_name} con IA..."):
            try:
                embeddings = HuggingFaceEmbeddings()
                vectordb_PEI = Chroma(persist_directory="./vector_db_PEI", embedding_function=embeddings)
                vectordb_guia = Chroma(persist_directory="./vector_db_guia", embedding_function=embeddings)
                vectordb_list = [vectordb_PEI, vectordb_guia]
                improved_text, context = mejorar_texto_con_IA(var_name, lineas_sentido, vectordb_list, nombre_curso)
                st.session_state['improved_' + var_name] = improved_text
                st.session_state['context_' + var_name] = context
                st.rerun()
            except Exception as e:
                st.warning(e)
                pass
    if 'context_' + var_name in st.session_state:
        st.warning("Fragmentos del PEI-UdeA usados por la IA")
        st.info(st.session_state['context_' + var_name])

    # 'derrotero_preguntas'
    var_name = 'derrotero_preguntas'
    var_value = st.session_state.get('improved_' + var_name, st.session_state.get(var_name, ''))
    derrotero_preguntas = st.text_area(
        "Derrotero de preguntas",
        value=var_value,
        key=var_name
    )
    if st.button("Mejorar con IA✨", key='btn_' + var_name):
        with st.spinner(f"Mejorando {var_name} con IA..."):
            try:
                embeddings = HuggingFaceEmbeddings()
                vectordb_PEI = Chroma(persist_directory="./vector_db_PEI", embedding_function=embeddings)
                vectordb_guia = Chroma(persist_directory="./vector_db_guia", embedding_function=embeddings)
                vectordb_list = [vectordb_PEI, vectordb_guia]
                improved_text, context = mejorar_texto_con_IA(var_name, derrotero_preguntas, vectordb_list, nombre_curso)
                st.session_state['improved_' + var_name] = improved_text
                st.session_state['context_' + var_name] = context
                st.rerun()
            except Exception as e:
                st.warning(e)
                pass
    if 'context_' + var_name in st.session_state:
        st.warning("Fragmentos del PEI-UdeA usados por la IA")
        st.info(st.session_state['context_' + var_name])
                                 
# For 'descripcion_conocimientos'
with st.expander("DESCRIPCIÓN DE LOS CONOCIMIENTOS Y/O SABERES"):
    var_name = 'descripcion_conocimientos'
    var_value = st.session_state.get('improved_' + var_name, st.session_state.get(var_name, ''))
    descripcion_conocimientos = st.text_area(
        "Explicitar los ejes problémicos, saberes, proyectos, contenidos o temas que se abordan en el desarrollo del curso.",
        value=var_value,
        key=var_name
    )
    if st.button("Mejorar con IA✨", key='btn_' + var_name):
        with st.spinner(f"Mejorando {var_name} con IA..."):
            try:
                embeddings = HuggingFaceEmbeddings()
                vectordb_PEI = Chroma(persist_directory="./vector_db_PEI", embedding_function=embeddings)
                vectordb_guia = Chroma(persist_directory="./vector_db_guia", embedding_function=embeddings)
                vectordb_list = [vectordb_PEI, vectordb_guia]
                improved_text, context = mejorar_texto_con_IA(var_name, descripcion_conocimientos, vectordb_list, nombre_curso)
                st.session_state['improved_' + var_name] = improved_text
                st.session_state['context_' + var_name] = context
                st.rerun()
            except Exception as e:
                st.warning(e)
                pass
    if 'context_' + var_name in st.session_state:
        st.warning("Fragmentos del PEI-UdeA usados por la IA")
        st.info(st.session_state['context_' + var_name])

# For 'metodologia' and others in "METODOLOGÍA"
with st.expander("METODOLOGÍA"):
    # Existing code for 'estrategias_didacticas' and others...
    estrategias_didacticas_options = [
        "Aprendizaje Basado en Problemas (ABP)",
        "Aprendizaje Basado en Proyectos (ABP)",
        "Aprendizaje invertido",
        "Aprendizaje Basado en Retos (ABR)",
        "Estudio de caso",
        "Aprendizaje entre pares",
        "Clase magistral",
        "Salida de campo",
        "Taller",
        "Otra(s)"
    ]
    estrategias_didacticas = st.multiselect("Estrategias didácticas:", estrategias_didacticas_options, default=st.session_state.get('estrategias_didacticas', []), key='estrategias_didacticas')
    if "Otra(s)" in estrategias_didacticas:
        estrategias_otras = st.text_input("Escriba el nombre de la estrategia.", value=st.session_state.get('estrategias_otras', ''), key='estrategias_otras')
    else:
        estrategias_otras = ''

    # 'metodologia'
    var_name = 'metodologia'
    var_value = st.session_state.get('improved_' + var_name, st.session_state.get(var_name, ''))
    metodologia = st.text_area(
        "Describa brevemente la metodología(s) utilizada(s).",
        value=var_value,
        key=var_name
    )
    if st.button("Mejorar con IA✨", key='btn_' + var_name):
        with st.spinner(f"Mejorando {var_name} con IA..."):
            try:
                embeddings = HuggingFaceEmbeddings()
                vectordb_PEI = Chroma(persist_directory="./vector_db_PEI", embedding_function=embeddings)
                vectordb_guia = Chroma(persist_directory="./vector_db_guia", embedding_function=embeddings)
                vectordb_list = [vectordb_PEI, vectordb_guia]
                improved_text, context = mejorar_texto_con_IA(var_name, metodologia, vectordb_list, nombre_curso)
                st.session_state['improved_' + var_name] = improved_text
                st.session_state['context_' + var_name] = context
                st.rerun()
            except Exception as e:
                st.warning(e)
                pass
    if 'context_' + var_name in st.session_state:
        st.warning("Fragmentos del PEI-UdeA usados por la IA")
        st.info(st.session_state['context_' + var_name])

    # 'medios_recursos'
    var_name = 'medios_recursos'
    var_value = st.session_state.get('improved_' + var_name, st.session_state.get(var_name, ''))
    medios_recursos = st.text_area(
        "Medios y recursos didácticos:",
        value=var_value,
        key=var_name
    )
    if st.button("Mejorar con IA✨", key='btn_' + var_name):
        with st.spinner(f"Mejorando {var_name} con IA..."):
            try:
                embeddings = HuggingFaceEmbeddings()
                vectordb_PEI = Chroma(persist_directory="./vector_db_PEI", embedding_function=embeddings)
                vectordb_guia = Chroma(persist_directory="./vector_db_guia", embedding_function=embeddings)
                vectordb_list = [vectordb_PEI, vectordb_guia]
                improved_text, context = mejorar_texto_con_IA(var_name, medios_recursos, vectordb_list, nombre_curso)
                st.session_state['improved_' + var_name] = improved_text
                st.session_state['context_' + var_name] = context
                st.rerun()
            except Exception as e:
                st.warning(e)
                pass
    if 'context_' + var_name in st.session_state:
        st.warning("Fragmentos del PEI-UdeA usados por la IA")
        st.info(st.session_state['context_' + var_name])

    # 'formas_interaccion'
    var_name = 'formas_interaccion'
    var_value = st.session_state.get('improved_' + var_name, st.session_state.get(var_name, ''))
    formas_interaccion = st.text_area(
        "Formas de interacción en los ambientes de aprendizaje y de acompañamiento del trabajo independiente del estudiante:",
        value=var_value,
        key=var_name
    )
    if st.button("Mejorar con IA✨", key='btn_' + var_name):
        with st.spinner(f"Mejorando {var_name} con IA..."):
            try:
                embeddings = HuggingFaceEmbeddings()
                vectordb_PEI = Chroma(persist_directory="./vector_db_PEI", embedding_function=embeddings)
                vectordb_guia = Chroma(persist_directory="./vector_db_guia", embedding_function=embeddings)
                vectordb_list = [vectordb_PEI, vectordb_guia]
                improved_text, context = mejorar_texto_con_IA(var_name, formas_interaccion, vectordb_list, nombre_curso)
                st.session_state['improved_' + var_name] = improved_text
                st.session_state['context_' + var_name] = context
                st.rerun()
            except Exception as e:
                st.warning(e)
                pass
    if 'context_' + var_name in st.session_state:
        st.warning("Fragmentos del PEI-UdeA usados por la IA")
        st.info(st.session_state['context_' + var_name])

    # 'estrategias_internacionalizacion'
    var_name = 'estrategias_internacionalizacion'
    var_value = st.session_state.get('improved_' + var_name, st.session_state.get(var_name, ''))
    estrategias_internacionalizacion = st.text_area(
        "Estrategias de internacionalización del currículo que se desarrollan para cumplir con las intencionalidades formativas del microcurrículo:",
        value=var_value,
        key=var_name
    )
    if st.button("Mejorar con IA✨", key='btn_' + var_name):
        with st.spinner(f"Mejorando {var_name} con IA..."):
            try:
                embeddings = HuggingFaceEmbeddings()
                vectordb_PEI = Chroma(persist_directory="./vector_db_PEI", embedding_function=embeddings)
                vectordb_guia = Chroma(persist_directory="./vector_db_guia", embedding_function=embeddings)
                vectordb_list = [vectordb_PEI, vectordb_guia]
                improved_text, context = mejorar_texto_con_IA(var_name, estrategias_internacionalizacion, vectordb_list, nombre_curso)
                st.session_state['improved_' + var_name] = improved_text
                st.session_state['context_' + var_name] = context
                st.rerun()
            except Exception as e:
                st.warning(e)
                pass
    if 'context_' + var_name in st.session_state:
        st.warning("Fragmentos del PEI-UdeA usados por la IA")
        st.info(st.session_state['context_' + var_name])

    # 'estrategias_diversidad'
    var_name = 'estrategias_diversidad'
    var_value = st.session_state.get('improved_' + var_name, st.session_state.get(var_name, ''))
    estrategias_diversidad = st.text_area(
        "Estrategias para abordar o visibilizar la diversidad desde la perspectiva de género, el enfoque diferencial o el enfoque intercultural:",
        value=var_value,
        key=var_name
    )
    if st.button("Mejorar con IA✨", key='btn_' + var_name):
        with st.spinner(f"Mejorando {var_name} con IA..."):
            try:
                embeddings = HuggingFaceEmbeddings()
                vectordb_PEI = Chroma(persist_directory="./vector_db_PEI", embedding_function=embeddings)
                vectordb_guia = Chroma(persist_directory="./vector_db_guia", embedding_function=embeddings)
                vectordb_list = [vectordb_PEI, vectordb_guia]
                improved_text, context = mejorar_texto_con_IA(var_name, estrategias_diversidad, vectordb_list, nombre_curso)
                st.session_state['improved_' + var_name] = improved_text
                st.session_state['context_' + var_name] = context
                st.rerun()
            except Exception as e:
                st.warning(e)
                pass
    if 'context_' + var_name in st.session_state:
        st.warning("Fragmentos del PEI-UdeA usados por la IA")
        st.info(st.session_state['context_' + var_name])

# For 'concepcion_evaluacion'
with st.expander("EVALUACIÓN"):
    # Variable 'concepcion_evaluacion'
    var_name = 'concepcion_evaluacion'
    var_value = st.session_state.get('improved_' + var_name, st.session_state.get(var_name, ''))
    concepcion_evaluacion = st.text_area(
        "Concepción de evaluación, modalidades y estrategias a través de las cuales se va a orientar.",
        value=var_value,
        key=var_name
    )
    if st.button("Mejorar con IA✨", key='btn_' + var_name):
        with st.spinner(f'Mejorando {var_name} con IA...'):
            try:
                embeddings = HuggingFaceEmbeddings()
                vectordb_PEI = Chroma(persist_directory="./vector_db_PEI", embedding_function=embeddings)
                vectordb_guia = Chroma(persist_directory="./vector_db_guia", embedding_function=embeddings)
                vectordb_list = [vectordb_PEI, vectordb_guia]
                improved_text, context = mejorar_texto_con_IA(var_name, concepcion_evaluacion, vectordb_list, nombre_curso)
                st.session_state['improved_' + var_name] = improved_text
                st.session_state['context_' + var_name] = context
                st.rerun()
            except Exception as e:
                 st.warning(e)
                 pass
    if 'context_' + var_name in st.session_state:
        st.warning("Fragmentos del PEI-UdeA usados por la IA")
        st.info(st.session_state['context_' + var_name])

    # Variable 'procesos_aprendizaje'
    var_name = 'procesos_aprendizaje'
    var_value = st.session_state.get('improved_' + var_name, st.session_state.get(var_name, ''))
    procesos_aprendizaje = st.text_area(
        "Procesos y resultados de aprendizaje del Programa Académico que se abordan en el curso.",
        value=var_value,
        key=var_name
    )
    if st.button("Mejorar con IA✨", key='btn_' + var_name):
        with st.spinner(f'Mejorando {var_name} con IA...'):
            try:
                embeddings = HuggingFaceEmbeddings()
                vectordb_PEI = Chroma(persist_directory="./vector_db_PEI", embedding_function=embeddings)
                vectordb_guia = Chroma(persist_directory="./vector_db_guia", embedding_function=embeddings)
                vectordb_list = [vectordb_PEI, vectordb_guia]
                improved_text, context = mejorar_texto_con_IA(var_name, procesos_aprendizaje, vectordb_list, nombre_curso)
                st.session_state['improved_' + var_name] = improved_text
                st.session_state['context_' + var_name] = context
                st.rerun()
            except Exception as e:
                 st.warning(e)
                 pass
    if 'context_' + var_name in st.session_state:
        st.warning("Fragmentos del PEI-UdeA usados por la IA")
        st.info(st.session_state['context_' + var_name])

    # Variable 'momentos_evaluacion'
    var_name = 'momentos_evaluacion'
    var_value = st.session_state.get('improved_' + var_name, st.session_state.get(var_name, ''))
    momentos_evaluacion = st.text_area(
        "Momentos y/o productos de la evaluación del curso y sus respectivos porcentajes.",
        value=var_value,
        key=var_name
    )
    if st.button("Mejorar con IA✨", key='btn_' + var_name):
        with st.spinner(f'Mejorando {var_name} con IA...'):
            try:
                embeddings = HuggingFaceEmbeddings()
                vectordb_PEI = Chroma(persist_directory="./vector_db_PEI", embedding_function=embeddings)
                vectordb_guia = Chroma(persist_directory="./vector_db_guia", embedding_function=embeddings)
                vectordb_list = [vectordb_PEI, vectordb_guia]
                improved_text, context = mejorar_texto_con_IA(var_name, momentos_evaluacion, vectordb_list, nombre_curso)
                st.session_state['improved_' + var_name] = improved_text
                st.session_state['context_' + var_name] = context
                st.rerun()
            except Exception as e:
                 st.warning(e)
                 pass
    if 'context_' + var_name in st.session_state:
        st.warning("Fragmentos del PEI-UdeA usados por la IA")
        st.info(st.session_state['context_' + var_name])
                                 

with st.expander("BIBLIOGRAFÍA Y OTRAS FUENTES"):
        bibliografia = st.text_area("Incluir solo la bibliografía que se requiere para el desarrollo del curso.", value=st.session_state.get('bibliografia', ''), key='bibliografia')
with st.expander("COMUNIDAD ACADÉMICA QUE PARTICIPÓ EN LA ELABORACIÓN DEL MICROCURRÍCULO"):
        nombres_apellidos = st.text_area("Nombres y apellidos:", value=st.session_state.get('nombres_apellidos', ''), key='nombres_apellidos')
        unidad_academica_participantes = st.text_area("Unidad académica:", value=st.session_state.get('unidad_academica_participantes', ''), key='unidad_academica_participantes')
        formacion_academica = st.text_area("Formación académica:", value=st.session_state.get('formacion_academica', ''), key='formacion_academica')
    


if st.button("Descargar Microcurriculum"):
    variables = st.session_state.to_dict()
    variables['estrategias_didacticas'] = ', '.join(variables.get('estrategias_didacticas', []))
    if "Otra(s)" in variables.get('estrategias_didacticas', '') and variables.get('estrategias_otras', ''):
        variables['estrategias_didacticas'] += f", {variables['estrategias_otras']}"
    doc_buffer = crear_documento(variables['nombre_curso'], variables)
    st.download_button(
        label="Descargar Microcurriculum",
        data=doc_buffer,
        file_name=f"Microcurriculum_{variables['nombre_curso']}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

